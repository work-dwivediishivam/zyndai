"""
File processing utilities for extracting content from various file formats.
Supports PDF, DOCX, Excel, and images.
"""
import io
import base64

from PIL import Image
from docx import Document
from PyPDF2 import PdfReader
from typing import Optional, Tuple
from openpyxl import load_workbook


class FileProcessor:
    """Process various file formats and extract content"""
    
    # Maximum characters per chunk to avoid context overflow
    MAX_CHUNK_SIZE = 50000  # ~12k tokens
    
    @staticmethod
    def process_file(filename: str, content: bytes, content_type: str) -> Tuple[str, Optional[str]]:
        """
        Process a file and extract its content.
        
        Args:
            filename: Name of the file
            content: Raw file bytes
            content_type: MIME type of the file
            
        Returns:
            Tuple of (base64_content, extracted_text)
        """
        base64_content = base64.b64encode(content).decode('utf-8')
        extracted_text = None
        
        try:
            if content_type == 'application/pdf':
                extracted_text = FileProcessor._extract_pdf(content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                extracted_text = FileProcessor._extract_docx(content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
                extracted_text = FileProcessor._extract_excel(content)
            elif content_type.startswith('image/'):
                # For images, we'll send the base64 directly to Gemini's vision model
                extracted_text = f"[Image: {filename}]"
        except Exception as e:
            print(f"Error processing file {filename}: {e}")
            extracted_text = f"[Error processing {filename}: {str(e)}]"
        
        return base64_content, extracted_text
    
    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        """Extract text from PDF"""
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            except Exception as e:
                text_parts.append(f"[Error extracting page {page_num + 1}: {str(e)}]")
        
        full_text = "\n\n".join(text_parts)
        
        # Chunk if too large
        if len(full_text) > FileProcessor.MAX_CHUNK_SIZE:
            full_text = FileProcessor._chunk_text(full_text, "PDF")
        
        return full_text
    
    @staticmethod
    def _extract_docx(content: bytes) -> str:
        """Extract text from DOCX"""
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        
        # Chunk if too large
        if len(full_text) > FileProcessor.MAX_CHUNK_SIZE:
            full_text = FileProcessor._chunk_text(full_text, "DOCX")
        
        return full_text
    
    @staticmethod
    def _extract_excel(content: bytes) -> str:
        """Extract data from Excel file"""
        excel_file = io.BytesIO(content)
        workbook = load_workbook(excel_file, read_only=True, data_only=True)
        
        text_parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===\n")
            
            rows_data = []
            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None for cell in row):
                    row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    rows_data.append(row_str)
            
            text_parts.append("\n".join(rows_data))
        
        full_text = "\n\n".join(text_parts)
        
        # Chunk if too large
        if len(full_text) > FileProcessor.MAX_CHUNK_SIZE:
            full_text = FileProcessor._chunk_text(full_text, "Excel")
        
        return full_text
    
    @staticmethod
    def _chunk_text(text: str, file_type: str) -> str:
        """
        Intelligently chunk large text to fit within context limits.
        Returns a summary or truncated version.
        """
        chunks = []
        current_chunk = []
        current_size = 0
        
        paragraphs = text.split('\n\n')
        
        for para in paragraphs:
            para_size = len(para)
            if current_size + para_size > FileProcessor.MAX_CHUNK_SIZE:
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                current_chunk = [para]
                current_size = para_size
            else:
                current_chunk.append(para)
                current_size += para_size
        
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        # If we have multiple chunks, take the first chunk and add a note
        if len(chunks) > 1:
            return f"{chunks[0]}\n\n[Note: This {file_type} file is very large. Showing first ~50,000 characters. Total chunks: {len(chunks)}]"
        
        return chunks[0] if chunks else text[:FileProcessor.MAX_CHUNK_SIZE]
    
    @staticmethod
    def validate_file_type(content_type: str) -> bool:
        """Check if file type is supported"""
        supported_types = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
            'image/jpeg',
            'image/png',
            'image/gif',
            'image/webp',
        ]
        return content_type in supported_types or content_type.startswith('image/')
